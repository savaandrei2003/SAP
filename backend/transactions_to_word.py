import openpyxl
import docx

# Define variable to load the dataframe
dataframe = openpyxl.load_workbook("transactions.xlsx")

# Define variable to read sheet
dataframe1 = dataframe.active

first_client_expenses = []
second_client_expenses = []
first_ok = 0
second_ok = 0

for row in range(0, dataframe1.max_row):
    for index, col in enumerate(dataframe1.iter_cols(0, dataframe1.max_column)):
        print("ok")
        if index == 0 and col[row].value == 14191:
            first_ok = 1
        elif index == 0:
            first_ok = 0
        if index == 0 and col[row].value == 16794:
            second_ok = 1
        elif index == 0:
            second_ok = 0

        if index == 3 and first_ok == 1:
            first_client_expenses.append(col[row].value)

        if index == 3 and second_ok == 1:
            second_client_expenses.append(col[row].value)


print(first_client_expenses)
print(second_client_expenses)

doc = docx.Document()
doc.add_paragraph(first_client_expenses)
doc.add_paragraph(second_client_expenses)
doc.save("clients_expenses.docx")
